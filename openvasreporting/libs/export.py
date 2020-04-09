# -*- coding: utf-8 -*-
#
#
# Project name: OpenVAS Reporting: A tool to convert OpenVAS XML reports into Excel files.
# Project URL: https://github.com/TheGroundZero/openvasreporting

import re
from collections import Counter

from .config import Config
from .parsed_data import Vulnerability

# DEBUG
import sys
import logging
#logging.basicConfig(stream=sys.stderr, level=logging.DEBUG,
#                    format="%(asctime)s | %(levelname)s | %(name)s | %(message)s", datefmt="%Y-%m-%d %H:%M:%S")
logging.basicConfig(stream=sys.stderr, level=logging.ERROR,
                    format="%(asctime)s | %(levelname)s | %(name)s | %(message)s", datefmt="%Y-%m-%d %H:%M:%S")


def exporters():
    """
    Enum-like instance containing references to correct exporter function

    > exporters()[key](param[s])

    :return: Pointer to exporter function
    """
    return {
        'xlsx': export_to_excel,
        'docx': export_to_word,
        'csv': export_to_csv
    }


def _get_collections(vuln_info):
    """
    Sort vulnerability list info according to CVSS (desc) and Name (asc).
    Provide collections to be used in export.

    :param vuln_info: Vulnerability list info
    :type vuln_info: list(Vulnerability)

    :return: vuln_info, vuln_levels, vuln_host_by_level, vuln_by_family
    :rtype vuln_info: list(Vulnerability)
    :rtype vuln_levels: Counter
    :rtype vuln_host_by_level: Counter
    :rtype vuln_by_family: Counter
    """
    vuln_info.sort(key=lambda key: key.name)
    vuln_info.sort(key=lambda key: key.cvss, reverse=True)
    vuln_levels = Counter()
    vuln_host_by_level = Counter()
    vuln_by_family = Counter()
    # collect host names
    vuln_hostcount_by_level =[[] for _ in range(5)]
    level_choices = {'critical': 0, 'high': 1, 'medium': 2, 'low': 3, 'none': 4}

    for i, vuln in enumerate(vuln_info, 1):
        vuln_levels[vuln.level.lower()] += 1
        # add host names to list so we count unquie hosts per level
        level_index = level_choices.get(vuln.level.lower())

        for i, (host, port) in enumerate(vuln.hosts, 1):    
            if host.ip not in vuln_hostcount_by_level[level_index]:
                vuln_hostcount_by_level[level_index].append(host.ip)       

        vuln_by_family[vuln.family] += 1

    # now count hosts per level and return
    for level in Config.levels().values():
        vuln_host_by_level[level] = len((vuln_hostcount_by_level[level_choices.get(level.lower())]))

    return vuln_info, vuln_levels, vuln_host_by_level, vuln_by_family


def export_to_excel(vuln_info, template=None, output_file='openvas_report.xlsx'):
    """
    Export vulnerabilities info in an Excel file.

    :param vuln_info: Vulnerability list info
    :type vuln_info: list(Vulnerability)
    :param template: Not supported in xlsx-output
    :type template: NoneType

    :param output_file: Filename of the Excel file
    :type output_file: str

    :raises: TypeError, NotImplementedError
    """

    import xlsxwriter

    if not isinstance(vuln_info, list):
        raise TypeError("Expected list, got '{}' instead".format(type(vuln_info)))
    else:
        for x in vuln_info:
            if not isinstance(x, Vulnerability):
                raise TypeError("Expected Vulnerability, got '{}' instead".format(type(x)))
    if not isinstance(output_file, str):
        raise TypeError("Expected str, got '{}' instead".format(type(output_file)))
    else:
        if not output_file:
            raise ValueError("output_file must have a valid name.")
    # if template is not None:
    #     raise NotImplementedError("Use of template is not supported in XSLX-output.")

    vuln_info, vuln_levels, vuln_host_by_level, vuln_by_family = _get_collections(vuln_info)

    # ====================
    # FUNCTIONS
    # ====================
    def __row_height(text, width):
        return (max((len(text) // width), text.count('\n')) + 1) * 15

    workbook = xlsxwriter.Workbook(output_file)

    workbook.set_properties({
        'title': output_file,
        'subject': 'OpenVAS report',
        'author': 'TheGroundZero',
        'category': 'report',
        'keywords': 'OpenVAS, report',
        'comments': 'TheGroundZero (https://github.com/TheGroundZero)'})

    # ====================
    # FORMATTING
    # ====================
    workbook.formats[0].set_font_name('Tahoma')

    format_sheet_title_content = workbook.add_format({'font_name': 'Tahoma', 'font_size': 12,
                                                      'font_color': Config.colors()['blue'], 'bold': True,
                                                      'align': 'center', 'valign': 'vcenter', 'border': 1})
    format_table_titles = workbook.add_format({'font_name': 'Tahoma', 'font_size': 11,
                                               'font_color': 'white', 'bold': True,
                                               'align': 'center', 'valign': 'vcenter',
                                               'border': 1,
                                               'bg_color': Config.colors()['blue']})
    format_table_cells = workbook.add_format({'font_name': 'Tahoma', 'font_size': 10,
                                              'align': 'left', 'valign': 'top',
                                              'border': 1, 'text_wrap': 1})
    format_align_center = workbook.add_format({'font_name': 'Tahoma', 'font_size': 10,
                                               'align': 'center', 'valign': 'top'})
    format_align_border = workbook.add_format({'font_name': 'Tahoma', 'font_size': 10,
                                               'align': 'center', 'valign': 'top',
                                               'border': 1, 'text_wrap': 1})
    format_toc = {
        'critical': workbook.add_format({'font_name': 'Tahoma', 'font_size': 10, 'font_color': 'white',
                                         'align': 'center', 'valign': 'top',
                                         'border': 1,
                                         'bg_color': Config.colors()['critical']}),
        'high': workbook.add_format({'font_name': 'Tahoma', 'font_size': 10, 'font_color': 'white',
                                     'align': 'center', 'valign': 'top',
                                     'border': 1, 'bg_color': Config.colors()['high']}),
        'medium': workbook.add_format({'font_name': 'Tahoma', 'font_size': 10, 'font_color': 'white',
                                       'align': 'center', 'valign': 'top',
                                       'border': 1, 'bg_color': Config.colors()['medium']}),
        'low': workbook.add_format({'font_name': 'Tahoma', 'font_size': 10, 'font_color': 'white',
                                    'align': 'center', 'valign': 'top',
                                    'border': 1, 'bg_color': Config.colors()['low']}),
        'none': workbook.add_format({'font_name': 'Tahoma', 'font_size': 10, 'font_color': 'white',
                                     'align': 'center', 'valign': 'top',
                                     'border': 1, 'bg_color': Config.colors()['none']})
    }

    # ====================
    # SUMMARY SHEET
    # ====================
    sheet_name = "Summary"
    ws_sum = workbook.add_worksheet(sheet_name)
    ws_sum.set_tab_color(Config.colors()['blue'])

    ws_sum.set_column("A:A", 7, format_align_center)
    ws_sum.set_column("B:B", 25, format_align_center)
    ws_sum.set_column("C:C", 24, format_align_center)
    ws_sum.set_column("D:D", 20, format_align_center)
    ws_sum.set_column("E:E", 7, format_align_center)

    # --------------------
    # VULN SUMMARY
    # --------------------
    ws_sum.merge_range("B2:D2", "VULNERABILITY SUMMARY", format_sheet_title_content)
    ws_sum.write("B3", "Threat", format_table_titles)
    ws_sum.write("C3", "Vulns number", format_table_titles)
    ws_sum.write("D3", "Affected hosts", format_table_titles)

    for i, level in enumerate(Config.levels().values(), 4):
        ws_sum.write("B{}".format(i), level.capitalize(), format_sheet_title_content)
        ws_sum.write("C{}".format(i), vuln_levels[level], format_align_border)
        ws_sum.write("D{}".format(i), vuln_host_by_level[level], format_align_border)

    ws_sum.write("B9", "Total", format_table_titles)
    ws_sum.write_formula("C9", "=SUM($C$4:$C$8)", format_table_titles)
    ws_sum.write_formula("D9", "=SUM($D$4:$D$8)", format_table_titles)

    # --------------------
    # CHART
    # --------------------
    chart_vulns_summary = workbook.add_chart({'type': 'pie'})
    chart_vulns_summary.add_series({
        'name': 'vulnerability summary by affected hosts',
        'categories': '={}!B4:B8'.format(sheet_name),
        'values': '={}!D4:D8'.format(sheet_name),
        'data_labels': {'value': True, 'position': 'outside_end', 'leader_lines': True, 'font': {'name': 'Tahoma'}},
        'points': [
            {'fill': {'color': Config.colors()['critical']}},
            {'fill': {'color': Config.colors()['high']}},
            {'fill': {'color': Config.colors()['medium']}},
            {'fill': {'color': Config.colors()['low']}},
            {'fill': {'color': Config.colors()['none']}},
        ],
    })
    chart_vulns_summary.set_title({'name': 'Vulnerability summary', 'overlay': False, 'name_font': {'name': 'Tahoma'}})
    chart_vulns_summary.set_size({'width': 500, 'height': 300})
    chart_vulns_summary.set_legend({'position': 'right', 'font': {'name': 'Tahoma'}})
    ws_sum.insert_chart("F2", chart_vulns_summary)

    # --------------------
    # VULN BY FAMILY
    # --------------------
    ws_sum.merge_range("B19:C19", "VULNERABILITIES BY FAMILY", format_sheet_title_content)
    ws_sum.write("B20", "family", format_table_titles)
    ws_sum.write("C20", "vulns number", format_table_titles)

    last = 21
    for i, (family, number) in enumerate(iter(vuln_by_family.items()), last):
        ws_sum.write("B{}".format(i), family, format_align_border)
        ws_sum.write("C{}".format(i), number, format_align_border)
        last = i

    ws_sum.write("B{}".format(str(last + 1)), "Total", format_table_titles)
    ws_sum.write_formula("C{}".format(str(last + 1)), "=SUM($C$21:$C${})".format(last), format_table_titles)

    # --------------------
    # CHART
    # --------------------
    chart_vulns_by_family = workbook.add_chart({'type': 'pie'})
    chart_vulns_by_family.add_series({
        'name': 'vulnerability summary by family',
        'categories': '={}!B21:B{}'.format(sheet_name, last),
        'values': '={}!C21:C{}'.format(sheet_name, last),
        'data_labels': {'value': True, 'position': 'best_fit', 'leader_lines': True, 'font': {'name': 'Tahoma'}},
    })
    chart_vulns_by_family.set_title({'name': 'Vulnerability by family', 'overlay': False,
                                     'name_font': {'name': 'Tahoma'}})
    chart_vulns_by_family.set_size({'width': 500, 'height': 500})
    chart_vulns_by_family.set_legend({'position': 'bottom', 'font': {'name': 'Tahoma'}})
    ws_sum.insert_chart("F19", chart_vulns_by_family)

    # ====================
    # TABLE OF CONTENTS
    # ====================
    sheet_name = "TOC"
    ws_toc = workbook.add_worksheet(sheet_name)
    ws_toc.set_tab_color(Config.colors()['blue'])

    ws_toc.set_column("A:A", 7)
    ws_toc.set_column("B:B", 5)
    ws_toc.set_column("C:C", 150)
    ws_toc.set_column("D:D", 15)
    ws_toc.set_column("E:E", 50)
    ws_toc.set_column("F:F", 7)

    ws_toc.merge_range("B2:E2", "TABLE OF CONTENTS", format_sheet_title_content)
    ws_toc.write("B3", "No.", format_table_titles)
    ws_toc.write("C3", "Vuln Title", format_table_titles)
    ws_toc.write("D3", "Level", format_table_titles)
    ws_toc.write("E3", "Hosts", format_table_titles)

    # ====================
    # VULN SHEETS
    # ====================
    for i, vuln in enumerate(vuln_info, 1):
        name = re.sub(r"[\[\]\\\'\"&@#():*?/]", "", vuln.name)
        if len(name) > 27:
            name = "{}..{}".format(name[0:15], name[-10:])
        name = "{:03X}_{}".format(i, name)
        ws_vuln = workbook.add_worksheet(name)
        ws_vuln.set_tab_color(Config.colors()[vuln.level.lower()])

        # --------------------
        # TABLE OF CONTENTS
        # --------------------
        ws_toc.write("B{}".format(i + 3), "{:03X}".format(i), format_table_cells)
        ws_toc.write_url("C{}".format(i + 3), "internal:'{}'!A1".format(name), format_table_cells, string=vuln.name)
        ws_toc.write("D{}".format(i + 3), "{:.1f} ({})".format(vuln.cvss, vuln.level.capitalize()),
                     format_toc[vuln.level])
        ws_toc.write("E{}".format(i + 3), "{}".format(', '.join([host.ip for host, _ in vuln.hosts])),
                     format_table_cells)
        ws_vuln.write_url("A1", "internal:'{}'!A{}".format(ws_toc.get_name(), i + 3), format_align_center,
                          string="<< TOC")
        ws_toc.set_row(i + 3, __row_height(name, 150), None)

        # --------------------
        # VULN INFO
        # --------------------
        ws_vuln.set_column("A:A", 7, format_align_center)
        ws_vuln.set_column("B:B", 20, format_align_center)
        ws_vuln.set_column("C:C", 20, format_align_center)
        ws_vuln.set_column("D:D", 50, format_align_center)
        ws_vuln.set_column("E:E", 15, format_align_center)
        ws_vuln.set_column("F:F", 15, format_align_center)
        ws_vuln.set_column("G:G", 20, format_align_center)
        ws_vuln.set_column("H:H", 7, format_align_center)
        content_width = 120

        ws_vuln.write('B2', "Title", format_table_titles)
        ws_vuln.merge_range("C2:G2", vuln.name, format_sheet_title_content)
        ws_vuln.set_row(1, __row_height(vuln.name, content_width), None)

        ws_vuln.write('B3', "Description", format_table_titles)
        ws_vuln.merge_range("C3:G3", vuln.description, format_table_cells)
        ws_vuln.set_row(2, __row_height(vuln.description, content_width), None)

        ws_vuln.write('B4', "Impact", format_table_titles)
        ws_vuln.merge_range("C4:G4", vuln.impact, format_table_cells)
        ws_vuln.set_row(3, __row_height(vuln.impact, content_width), None)

        ws_vuln.write('B5', "Recommendation", format_table_titles)
        ws_vuln.merge_range("C5:G5", vuln.solution, format_table_cells)
        ws_vuln.set_row(4, __row_height(vuln.solution, content_width), None)

        ws_vuln.write('B6', "Details", format_table_titles)
        ws_vuln.merge_range("C6:G6", vuln.insight, format_table_cells)
        ws_vuln.set_row(5, __row_height(vuln.insight, content_width), None)

        ws_vuln.write('B7', "CVEs", format_table_titles)
        cves = ", ".join(vuln.cves)
        cves = cves.upper() if cves != "" else "No CVE"
        ws_vuln.merge_range("C7:G7", cves, format_table_cells)
        ws_vuln.set_row(6, __row_height(cves, content_width), None)

        ws_vuln.write('B8', "CVSS", format_table_titles)
        cvss = float(vuln.cvss)
        if cvss >= 0.0:
            ws_vuln.merge_range("C8:G8", "{:.1f}".format(cvss), format_table_cells)
        else:
            ws_vuln.merge_range("C8:G8", "{}".format("No CVSS"), format_table_cells)

        ws_vuln.write('B9', "Level", format_table_titles)
        ws_vuln.merge_range("C9:G9", vuln.level.capitalize(), format_table_cells)

        ws_vuln.write('B10', "Family", format_table_titles)
        ws_vuln.merge_range("C10:G10", vuln.family, format_table_cells)

        ws_vuln.write('B11', "References", format_table_titles)
        ws_vuln.merge_range("C11:G11", vuln.references, format_table_cells)
        ws_vuln.set_row(10, __row_height(vuln.references, content_width), None)

        ws_vuln.write('C13', "IP", format_table_titles)
        ws_vuln.write('D13', "Host name", format_table_titles)
        ws_vuln.write('E13', "Port number", format_table_titles)
        ws_vuln.write('F13', "Port protocol", format_table_titles)
        ws_vuln.write('G13', "Port Result", format_table_titles)

        # --------------------
        # AFFECTED HOSTS
        # --------------------
        for j, (host, port) in enumerate(vuln.hosts, 14):

            ws_vuln.write("C{}".format(j), host.ip)
            ws_vuln.write("D{}".format(j), host.host_name if host.host_name else "-")

            if port:
                ws_vuln.write("E{}".format(j), "" if port.number == 0 else port.number)
                ws_vuln.write("F{}".format(j), port.protocol)
                ws_vuln.write("G{}".format(j), port.result, format_table_cells)
                ws_vuln.set_row(j, __row_height(port.result, content_width), None)
            else:
                ws_vuln.write("E{}".format(j), "No port info")

    workbook.close()


def export_to_word(vuln_info, template, output_file='openvas_report.docx'):
    """
    Export vulnerabilities info in a Word file.

    :param vuln_info: Vulnerability list info
    :type vuln_info: list(Vulnerability)

    :param output_file: Filename of the Excel file
    :type output_file: str
    
    :param template: Path to Docx template
    :type template: str

    :raises: TypeError
    """

    import matplotlib.pyplot as plt
    import numpy as np
    import tempfile
    import os
    import math

    from docx import Document
    from docx.oxml.shared import qn, OxmlElement
    from docx.oxml.ns import nsdecls
    from docx.oxml import parse_xml
    from docx.shared import Cm, Pt, Twips
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    from docx.enum.section import WD_ORIENT
    from docx.enum.style import WD_STYLE_TYPE
    from docx.enum.table import WD_ALIGN_VERTICAL
    from docx.shared import RGBColor

    if not isinstance(vuln_info, list):
        raise TypeError("Expected list, got '{}' instead".format(type(vuln_info)))
    else:
        for x in vuln_info:
            if not isinstance(x, Vulnerability):
                raise TypeError("Expected Vulnerability, got '{}' instead".format(type(x)))
    if not isinstance(output_file, str):
        raise TypeError("Expected str, got '{}' instead".format(type(output_file)))
    else:
        if not output_file:
            raise ValueError("output_file must have a valid name.")
    if template is not None:
        if not isinstance(template, str):
            raise TypeError("Expected str, got '{}' instead".format(type(template)))

    vuln_info, vuln_levels, vuln_host_by_level, vuln_by_family = _get_collections(vuln_info)

    # ====================
    # DOCUMENT PROPERTIES
    # ====================
    # Create new doc
    if template is None:
        document = Document()
        doc_section = document.sections[0]
        # Set A4 Format
        doc_section.page_width    = Cm(21.0)
        doc_section.page_height   = Cm(29.7)
        # Shrink margins almost to 0
        doc_section.left_margin   = Cm(1.5)
        doc_section.right_margin  = doc_section.left_margin
        doc_section.top_margin    = Cm(1.0)
        doc_section.bottom_margin = Cm(1.0)
        # Force portrait
        doc_section.orientation   = WD_ORIENT.PORTRAIT
    # use template
    else:
        document = Document(template)
        doc_section = document.sections[0]
    
    # Defining styles (if not defined already)
    # All used style will be custom, and with 'OR-' prefix.
    # In this way, the template can still define styles.
    doc_styles = document.styles
    
    # Base paragraph
    if 'OR-base' not in doc_styles:
        style_pr_base = doc_styles.add_style('OR-base', WD_STYLE_TYPE.PARAGRAPH)
        style_pr_base.font.name = 'Ubuntu'
        style_pr_base.font.size = Pt(8)
        style_pr_base.font.color.rgb                     = RGBColor.from_string('080808')
        style_pr_base.paragraph_format.left_indent       = Cm(0)
        style_pr_base.paragraph_format.right_indent      = Cm(0)
        style_pr_base.paragraph_format.first_line_indent = Cm(0)
        style_pr_base.paragraph_format.space_before      = Cm(0)
        style_pr_base.paragraph_format.space_after       = Cm(0)
        style_pr_base.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        style_pr_base.paragraph_format.widow_control     = True
        style_pr_base.paragraph_format.alignment         = WD_ALIGN_PARAGRAPH.JUSTIFY_LOW
    # Base Styles modification
    if 'OR-base_bold' not in doc_styles:
        style_pr_body = doc_styles.add_style('OR-base_bold', WD_STYLE_TYPE.PARAGRAPH)
        style_pr_body.base_style = doc_styles['OR-base']
        style_pr_body.font.bold  = True
    # Section headers
    if 'OR-Heading_base' not in doc_styles:
        style_pr_or_head_base = doc_styles.add_style('OR-Heading_base', WD_STYLE_TYPE.PARAGRAPH)
        style_pr_or_head_base.base_style                   = doc_styles['OR-base_bold']
        style_pr_or_head_base.font.color.rgb               = RGBColor.from_string('183868')
        style_pr_or_head_base.paragraph_format.space_after = Pt(4)
    # - Titles
    if 'OR-Title' not in doc_styles:
        style_pr_or_title = doc_styles.add_style('OR-Title', WD_STYLE_TYPE.PARAGRAPH)
        style_pr_or_title.base_style                 = doc_styles['OR-Heading_base']
        style_pr_or_title.font.size                  = Pt(36)
        style_pr_or_title.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        style_pr_or_head_base.paragraph_format.space_after = Pt(8)
    # - Headers
    if 'OR-Heading_1' not in doc_styles:
        style_pr_or_header = doc_styles.add_style('OR-Heading_1', WD_STYLE_TYPE.PARAGRAPH)
        style_pr_or_header.base_style = doc_styles['OR-Heading_base']
        style_pr_or_header.font.size  = Pt(20)
    if 'OR-Heading_2' not in doc_styles:
        style_pr_or_header = doc_styles.add_style('OR-Heading_2', WD_STYLE_TYPE.PARAGRAPH)
        style_pr_or_header.base_style = doc_styles['OR-Heading_base']
        style_pr_or_header.font.size  = Pt(16)
    if 'OR-Heading_3' not in doc_styles:
        style_pr_or_header = doc_styles.add_style('OR-Vuln_title', WD_STYLE_TYPE.PARAGRAPH)
        style_pr_or_header.base_style     = doc_styles['OR-Heading_base']
        style_pr_or_header.font.size      = Pt(12)
    # - Vulnerabilities Titles
    for name,rgb in Config.colors().items():
        if 'OR-Vuln_title_'+name not in doc_styles:
            style_pr_or_header = doc_styles.add_style('OR-Vuln_title_'+name, WD_STYLE_TYPE.PARAGRAPH)
            style_pr_or_header.base_style     = doc_styles['OR-Vuln_title']
            style_pr_or_header.font.color.rgb = RGBColor.from_string(rgb[1:])
    # - Host with vulnerabilities title
    if 'OR-Vuln_hosts' not in doc_styles:
        style_pr_or_header = doc_styles.add_style('OR-Vuln_hosts', WD_STYLE_TYPE.PARAGRAPH)
        style_pr_or_header.base_style     = doc_styles['OR-Heading_base']
        style_pr_or_header.font.size      = Pt(10)
    # TOC specific
    if 'OR-TOC_base' not in doc_styles:
        style_pr_or_toc_base = doc_styles.add_style('OR-TOC_base', WD_STYLE_TYPE.PARAGRAPH)
        style_pr_or_toc_base.base_style     = doc_styles['OR-base']
        style_pr_or_toc_base.font.color.rgb = RGBColor.from_string('183868')
    if 'OR-TOC_1' not in doc_styles:
        style_pr_or_toc = doc_styles.add_style('OR-TOC_1', WD_STYLE_TYPE.PARAGRAPH)
        style_pr_or_toc.base_style = doc_styles['OR-TOC_base']
        style_pr_or_toc.font.bold  = True
    if 'OR-TOC_2' not in doc_styles:
        style_pr_or_toc = doc_styles.add_style('OR-TOC_2', WD_STYLE_TYPE.PARAGRAPH)
        style_pr_or_toc.base_style = doc_styles['OR-TOC_base']
    if 'OR-TOC_3' not in doc_styles:
        style_pr_or_toc = doc_styles.add_style('OR-Toc_3', WD_STYLE_TYPE.PARAGRAPH)
        style_pr_or_toc.base_style = doc_styles['OR-TOC_base']
    if 'OR-TOC_4' not in doc_styles:
        style_pr_or_toc = doc_styles.add_style('OR-TOC_4', WD_STYLE_TYPE.PARAGRAPH)
        style_pr_or_toc.base_style = doc_styles['OR-TOC_base']
        style_pr_or_toc.font.italic  = True
    # Tables style
    # - Specific paragraph style to allow space before and after
    if 'OR-cell' not in doc_styles:
        style_pr_cell = doc_styles.add_style('OR-cell', WD_STYLE_TYPE.PARAGRAPH)
        style_pr_cell.base_style                    = doc_styles['OR-base']
        style_pr_cell.paragraph_format.space_before = Pt(1.5)
        style_pr_cell.paragraph_format.space_after  = Pt(1.5)
    if 'OR-cell_bold' not in doc_styles:
        style_pr_cell = doc_styles.add_style('OR-cell_bold', WD_STYLE_TYPE.PARAGRAPH)
        style_pr_cell.base_style = doc_styles['OR-cell']
        style_pr_cell.font.bold  = True
    
    # Clear all contents
    document._body.clear_content()
    
    # Set doc title
    doc_prop    = document.core_properties
    doc_title = 'OpenVAS/GreenBone Report'
    doc_prop.title = doc_title
    doc_prop.category = "Report"
    
    # Effective writeable width
    # If margins set are float, try to fix (issue in python-docx: expected an int)
    # In this case, they _should be_ in twentieths of a point, so
    # multiply Twips helper
    try:
        doc_section.left_margin
    except ValueError as e:
        fixed_margin = float(re.search(": '(.+?)'", str(e)).group(1))
        doc_section.left_margin = Twips(fixed_margin)
    try:
        doc_section.right_margin
    except ValueError as e:
        fixed_margin = float(re.search(": '(.+?)'", str(e)).group(1))
        doc_section.right_margin = Twips(fixed_margin)
    
    page_width = doc_section.page_width - ( doc_section.left_margin + doc_section.right_margin )
    
    ## Start actual document writing ##
    document.add_paragraph(doc_title, style='OR-Title')

    # ====================
    # TABLE OF CONTENTS
    # ====================
    # WARNING -Not working with LibreOffice
    document.add_paragraph('Table of Contents', style='OR-Heading_1')
    # keep the title as cover of the report
    document.add_page_break()
    
    par = document.add_paragraph(style='OR-base')
    run = par.add_run()
    fld_char = OxmlElement('w:fldChar')  # creates a new element
    fld_char.set(qn('w:fldCharType'), 'begin')  # sets attribute on element
    instr_text = OxmlElement('w:instrText')
    instr_text.set(qn('xml:space'), 'preserve')  # sets attribute on element
    instr_text.text = r'TOC \h \z \t "OR-TOC_1;1;OR-OR-TOC_1;2;OR-TOC_3;3;OR-TOC_4;3"'

    fld_char2 = OxmlElement('w:fldChar')
    fld_char2.set(qn('w:fldCharType'), 'separate')
    fld_char3 = OxmlElement('w:t')
    fld_char3.text = "# Right-click to update field. #"
    fld_char2.append(fld_char3)

    fld_char4 = OxmlElement('w:fldChar')
    fld_char4.set(qn('w:fldCharType'), 'end')

    r_element = run._r
    r_element.append(fld_char)
    r_element.append(instr_text)
    r_element.append(fld_char2)
    r_element.append(fld_char4)

    document.add_page_break()

    # ====================
    # MANAGEMENT SUMMARY
    # ====================
    document.add_paragraph('Management Summary', style='OR-Heading_1')
    document.add_paragraph('< TYPE YOUR MANAGEMENT SUMMARY HERE >', style='OR-base')
    document.add_page_break()

    # ====================
    # TECHNICAL FINDINGS
    # ====================
    document.add_paragraph('Technical Findings', style='OR-Heading_1')
    document.add_paragraph('The section below discusses the technical findings.', style='OR-base' )

    # --------------------
    # SUMMARY TABLE
    # --------------------
    document.add_paragraph('Summary', style='OR-Heading_2')

    colors_sum = []
    labels_sum = []
    vuln_sum = []
    aff_sum = []

    table_summary = document.add_table(rows=1, cols=3)
    
    # TABLE HEADERS
    # --------------------
    hdr_cells = table_summary.rows[0].cells
    hdr_cells[0].paragraphs[0].add_run('Risk level')
    hdr_cells[1].paragraphs[0].add_run('Vulns number')
    hdr_cells[2].paragraphs[0].add_run('Affected hosts')
    # FIELDS
    # --------------------
    # Provide data to table and charts
    for level in Config.levels().values():
        row_cells = table_summary.add_row().cells
        row_cells[0].text = level.capitalize()
        row_cells[1].text = str(vuln_levels[level])
        row_cells[2].text = str(vuln_host_by_level[level])
        colors_sum.append(Config.colors()[level])
        labels_sum.append(level)
        vuln_sum.append(vuln_levels[level])
        aff_sum.append(vuln_host_by_level[level])
    
    # Apply styles
    # --------------------
    for h in hdr_cells:
        h.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        for p in h.paragraphs:
            p.style = doc_styles['OR-cell_bold']
    
    for r in table_summary.rows[1:]:
        for c in r.cells:
            c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for p in c.paragraphs:
                p.style = doc_styles['OR-cell']

    # --------------------
    # CHART
    # --------------------
    fd, path = tempfile.mkstemp(suffix='.png')
    chart_dpi    = 144
    chart_height = Cm(8);
    par_chart = document.add_paragraph(style='OR-base')
    par_chart.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_chart = par_chart.add_run()

    bar_chart, bar_axis = plt.subplots(dpi=chart_dpi)
    bar_axis.set_title('Vulnerability summary by risk level', fontsize=10)

    pos = np.arange(len(labels_sum))
    width = 0.35

    bar_axis.set_xticks(pos)
    bar_axis.set_xticklabels(labels_sum)
    bar_chart.gca().spines['left'].set_visible(False)
    bar_chart.gca().spines['right'].set_visible(False)
    bar_chart.gca().spines['top'].set_visible(False)
    bar_chart.gca().spines['bottom'].set_position('zero')
    bar_axis.tick_params(top=False, bottom=True, left=False, right=False,
                    labelleft=False, labelbottom=True)
    bars_vuln = plt.bar(pos - width / 2, vuln_sum, width, align='center', label='Vulnerabilities',
                        color=colors_sum, edgecolor='black')
    bars_aff = plt.bar(pos + width / 2, aff_sum, width, align='center', label='Affected hosts',
                       color=colors_sum, edgecolor='black', hatch='//')
    for barcontainer in (bars_vuln, bars_aff):
        for bar in barcontainer:
            height = bar.get_height()
            bar_chart.gca().text(bar.get_x() + bar.get_width() / 2, bar.get_height() + 0.3, str(int(height)),
                           ha='center', color='black', fontsize=8)
    bar_chart.legend()

    bar_chart.savefig(path)

    # plt.show()  # DEBUG
    
    bar_height = chart_height
    run_chart.add_picture(path, height=bar_height)
    os.remove(path)

    pie_chart, pie_axis = plt.subplots(dpi=chart_dpi, subplot_kw=dict(aspect="equal"))
    pie_axis.set_title('Vulnerability by family', fontsize=10)
    
    values = list(vuln_by_family.values())
    pie, tx, autotexts = pie_axis.pie(values, labels=vuln_by_family.keys(), autopct='', textprops=dict(fontsize=8))
    for i, txt in enumerate(autotexts):
        txt.set_text('{}'.format(values[i]))
    pie_chart.savefig(path)

    # plt.show()  # DEBUG
    pie_height = chart_height
    run_chart.add_picture(path, height=pie_height)
    os.remove(path)

    # ====================
    # VULN PAGES
    # ====================
    cur_level = ""

    for i, vuln in enumerate(vuln_info, 1):
        # --------------------
        # GENERAL
        # --------------------
        level = vuln.level.lower()

        if level != cur_level:
            document.add_paragraph(
                level.capitalize(), style='OR-Heading_2').paragraph_format.page_break_before = True
            cur_level = level
        else:
            document.add_page_break()

        title = "[{}] {}".format(level.upper(), vuln.name)
        par = document.add_paragraph(title, style='OR-Vuln_title_'+vuln.level)
        
        table_vuln = document.add_table(rows=8, cols=3)
        table_vuln.autofit = False
        table_vuln.columns[0].width = Cm(0.35)
        table_vuln.columns[1].width = Cm(2.85)
        table_vuln.columns[-1].width = page_width
        for c in range (len(table_vuln.columns)-1):
            table_vuln.columns[-1].width -= table_vuln.columns[c].width

        # COLOR
        # --------------------
        col_cells = table_vuln.columns[0].cells
        col_cells[0].merge(col_cells[7])
        color_fill = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), Config.colors()[vuln.level][1:]))
        col_cells[0]._tc.get_or_add_tcPr().append(color_fill)

        # TABLE HEADERS
        # --------------------
        hdr_cells = table_vuln.columns[1].cells
        hdr_cells[0].paragraphs[0].add_run('Description')
        hdr_cells[1].paragraphs[0].add_run('Impact')
        hdr_cells[2].paragraphs[0].add_run('Recommendation')
        hdr_cells[3].paragraphs[0].add_run('Details')
        hdr_cells[4].paragraphs[0].add_run('CVSS')
        hdr_cells[5].paragraphs[0].add_run('CVEs')
        hdr_cells[6].paragraphs[0].add_run('Family')
        hdr_cells[7].paragraphs[0].add_run('References')

        # FIELDS
        # --------------------
        cves = ", ".join(vuln.cves)
        cves = cves.upper() if cves != "" else "No CVE"

        cvss = str(vuln.cvss) if vuln.cvss != -1.0 else "No CVSS"

        txt_cells = table_vuln.columns[2].cells
        txt_cells[0].text = vuln.description
        txt_cells[1].text = vuln.impact
        txt_cells[2].text = vuln.solution
        txt_cells[3].text = vuln.insight
        txt_cells[4].text = cvss
        txt_cells[5].text = cves
        txt_cells[6].text = vuln.family
        txt_cells[7].text = vuln.references
        for c in txt_cells:
            for p in c.paragraphs:
                p.style = doc_styles['OR-cell']
        
        # Apply styles
        # --------------------
        for h in hdr_cells:
            for p in h.paragraphs:
                p.style = doc_styles['OR-cell_bold']
        
        for c in txt_cells:
            for p in c.paragraphs:
                p.style = doc_styles['OR-cell']

        # VULN HOSTS
        # --------------------
        document.add_paragraph('Vulnerable hosts', style='OR-Vuln_hosts')

        # add coloumn for result per port and resize columns
        table_hosts = document.add_table(cols=5, rows=(len(vuln.hosts) + 1))

        table_hosts.columns[0].width  = Cm(2.8)
        table_hosts.columns[1].width  = Cm(3.0)
        table_hosts.columns[2].width  = Cm(2.0)
        table_hosts.columns[3].width  = Cm(2.0)
        table_hosts.columns[-1].width = page_width
        for c in range (len(table_hosts.columns)-1):
            table_hosts.columns[-1].width -= table_hosts.columns[c].width
        
        # TABLE HEADERS
        # --------------------
        hdr_cells = table_hosts.rows[0].cells
        hdr_cells[0].paragraphs[0].add_run('IP')
        hdr_cells[1].paragraphs[0].add_run('Host name')
        hdr_cells[2].paragraphs[0].add_run('Port number')
        hdr_cells[3].paragraphs[0].add_run('Port protocol')
        hdr_cells[4].paragraphs[0].add_run('Port result')
        # FIELDS
        # --------------------
        for j, (host, port) in enumerate(vuln.hosts, 1):
            cells = table_hosts.rows[j].cells
            cells[0].text = host.ip
            cells[1].text = host.host_name if host.host_name else "-"
            if port and port is not None:
                cells[2].text = "-" if port.number == 0 else str(port.number)
                cells[3].text = port.protocol
                cells[4].text = port.result
            else:
                cells[2].text = "No port info"
        
        # Apply styles
        # --------------------
        for h in hdr_cells:
            for p in h.paragraphs:
                p.style = doc_styles['OR-cell_bold']
        for r in table_hosts.rows[1:]:
            for c in r.cells:
                for p in c.paragraphs:
                    p.style = doc_styles['OR-cell']

    document.save(output_file)


def export_to_csv(vuln_info, template=None, output_file='openvas_report.csv'):
    """
    Export vulnerabilities info in a Comma Separated Values (csv) file

    :param vuln_info: Vulnerability list info
    :type vuln_info: list(Vulnerability)

    :param template: Not supported in csv-output
    :type template: NoneType

    :param output_file: Filename of the csv file
    :type output_file: str

    :raises: TypeError, NotImplementedError
    """

    import csv

    if not isinstance(vuln_info, list):
        raise TypeError("Expected list, got '{}' instead".format(type(vuln_info)))
    else:
        for x in vuln_info:
            if not isinstance(x, Vulnerability):
                raise TypeError("Expected Vulnerability, got '{}' instead".format(type(x)))
    if not isinstance(output_file, str):
        raise TypeError("Expected str, got '{}' instead".format(type(output_file)))
    else:
        if not output_file:
            raise ValueError("output_file must have a valid name.")
    # Make a warning, not an error
    if template is not None:
        print("WARNING: Use of template is not supported in CSV-output.", file=sys.stderr)

    vuln_info, _, _, _ = _get_collections(vuln_info)

    with open(output_file, 'w') as csvfile:
        fieldnames = ['hostname', 'ip', 'port', 'protocol',
                      'vulnerability', 'cvss', 'threat', 'family',
                      'description', 'detection', 'insight', 'impact', 'affected', 'solution', 'solution_type',
                      'vuln_id', 'cve', 'references']
        writer = csv.DictWriter(csvfile, dialect='excel', fieldnames=fieldnames)
        writer.writeheader()

        for vuln in vuln_info:
            for (host, port) in vuln.hosts:
                rowdata = {
                    'hostname': host.host_name,
                    'ip': host.ip,
                    'port': port.number,
                    'protocol': port.protocol,
                    'vulnerability': vuln.name,
                    'cvss': vuln.cvss,
                    'threat': vuln.level,
                    'family': vuln.family,
                    'description': vuln.description,
                    'detection': vuln.detect,
                    'insight': vuln.insight,
                    'impact': vuln.impact,
                    'affected': vuln.affected,
                    'solution': vuln.solution,
                    'solution_type': vuln.solution_type,
                    'vuln_id': vuln.vuln_id,
                    'cve': ' - '.join(vuln.cves),
                    'references': ' - '.join(vuln.references)
                }
                writer.writerow(rowdata)
